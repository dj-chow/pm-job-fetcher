#!/usr/bin/env python3
"""
Daily PM job digest: fetches new jobs, customizes resumes via Claude CLI,
checks for hallucinations, and emails top matches with .docx attachments.

Uses your existing Claude subscription via the 'claude -p' CLI command.
No separate API key required.

Usage:
    python daily_digest.py              # fetch new jobs + send email
    python daily_digest.py --test       # dry run (no email, no Claude calls)
    python daily_digest.py --today      # use today's already-fetched output

Configure settings.json before running:
    email_from, email_app_password, email_to
"""

import json
import os
import re
import sys
import datetime
import zipfile
import smtplib
import subprocess
import urllib.request
import urllib.error
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))


def _safe_print(text):
    """Print with ASCII fallback so Windows cp1252 consoles don't crash."""
    try:
        print(text)
    except UnicodeEncodeError:
        print(text.encode("ascii", errors="replace").decode("ascii"))


# Claude CLI — use .cmd wrapper on Windows so it resolves via npm
CLAUDE_CMD = ["claude.cmd"] if os.name == "nt" else ["claude"]
SETTINGS_FILE = os.path.join(SCRIPT_DIR, "settings.json")
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "output")
RESUMES_DIR = os.path.join(SCRIPT_DIR, "customized_resumes")
USER_AGENT = "Mozilla/5.0 (compatible; PMJobFetcher/1.0)"

# ── Scoring config for candidate's profile ───────────────────────────────────
# Adjust these weights to tune which jobs surface in your daily email
SCORE_WEIGHTS = {
    "senior_level":   ("title", ["senior", "staff", "principal", "head of", "group product", " lead"], 3),
    "too_senior":     ("title", ["vp ", "vice president", "chief"], -4),
    "too_junior":     ("title", ["associate pm", " apm ", "entry"], -5),
    "risk_domain":    ("title", ["risk", "compliance", "regulatory", "fraud", "financial reporting", "audit"], 8),
    "fintech_domain": ("title", ["payments", "payment", "card", "fintech", "financial"], 5),
    "platform_domain":("title", ["data", "platform", "enterprise", "b2b"], 3),
    "canada":         ("location", ["canada"], 5),
    "toronto":        ("location", ["toronto", "ontario", "mississauga", "vancouver"], 4),
    "remote":         ("location", ["remote"], 2),
    "fintech_cos":    ("company", ["affirm", "plaid", "stripe", "brex", "chime", "robinhood",
                                   "coinbase", "ramp", "wealthsimple", "koho", "wave", "nesto",
                                   "stackadapt", "lyft", "instacart", "gitlab"], 3),
}


# ── Settings ────────────────────────────────────────────────────────────────

def load_settings():
    if not os.path.exists(SETTINGS_FILE):
        print("ERROR: settings.json not found.")
        sys.exit(1)
    with open(SETTINGS_FILE) as f:
        return json.load(f)


def validate_settings(s, test_mode=False):
    missing = []
    if not s.get("email_from"):
        missing.append("email_from")
    if not s.get("email_app_password"):
        missing.append("email_app_password")
    if missing and not test_mode:
        print(f"ERROR: Missing settings: {', '.join(missing)}")
        print("Fill these in settings.json before running.")
        sys.exit(1)
    # Verify claude CLI is available
    result = subprocess.run(CLAUDE_CMD + ["--version"], capture_output=True, text=True)
    if result.returncode != 0 and not test_mode:
        print("ERROR: 'claude' CLI not found. Make sure Claude Code is installed and logged in.")
        sys.exit(1)
    return True


# ── Base resume loading ──────────────────────────────────────────────────────

def load_base_resume(settings):
    """Extract plain text from the .docx base resume."""
    path = settings.get("base_resume_path", "")
    if not path or not os.path.exists(path):
        print(f"WARNING: base_resume_path not found: {path}")
        return ""
    try:
        import zipfile as zf
        with zf.ZipFile(path) as z:
            with z.open("word/document.xml") as f:
                content = f.read().decode("utf-8", errors="replace")
                text = re.sub(r"<[^>]+>", " ", content)
                text = re.sub(r"\s+", " ", text).strip()
                return text
    except Exception as e:
        print(f"WARNING: Could not read base resume: {e}")
        return ""


# ── Skill loading ────────────────────────────────────────────────────────────

def load_skill_content(settings):
    """Extract the SKILL.md text from the .skill zip archive."""
    path = settings.get("skill_path", "")
    if not path or not os.path.exists(path):
        return ""
    try:
        with zipfile.ZipFile(path) as z:
            with z.open("resume-customizer/SKILL.md") as f:
                return f.read().decode("utf-8", errors="replace")
    except Exception as e:
        print(f"WARNING: Could not read skill file: {e}")
        return ""


# ── Job parsing ──────────────────────────────────────────────────────────────

def parse_jobs_from_file(filepath):
    """Parse the markdown output file into a list of job dicts."""
    jobs = []
    current_company = None
    try:
        with open(filepath, encoding="utf-8", errors="replace") as f:
            for line in f:
                if line.startswith("## "):
                    current_company = line.strip().replace("## ", "")
                elif "| " in line and "[Apply]" in line:
                    parts = [p.strip() for p in line.split("|")]
                    if len(parts) >= 4:
                        title = parts[1]
                        location = parts[2]
                        link_match = re.search(r"\[Apply\]\((.*?)\)", parts[3])
                        link = link_match.group(1) if link_match else ""
                        jobs.append({
                            "company": current_company or "",
                            "title": title,
                            "location": location,
                            "link": link,
                        })
    except Exception as e:
        print(f"ERROR parsing jobs file: {e}")
    return jobs


def find_latest_output():
    """Return path to today's output file, or the most recent one."""
    today = datetime.date.today().strftime("jobs_%Y-%m-%d.md")
    today_path = os.path.join(OUTPUT_DIR, today)
    if os.path.exists(today_path):
        return today_path
    # Find most recent
    files = sorted([
        f for f in os.listdir(OUTPUT_DIR) if f.startswith("jobs_") and f.endswith(".md")
    ], reverse=True)
    if files:
        return os.path.join(OUTPUT_DIR, files[0])
    return None


# ── Job scoring ──────────────────────────────────────────────────────────────

def score_job(job):
    score = 0
    for key, (field, keywords, weight) in SCORE_WEIGHTS.items():
        text = job.get(field, "").lower()
        if any(kw in text for kw in keywords):
            score += weight
    return score


# ── Location filter ───────────────────────────────────────────────────────
# Only surface jobs in Toronto, Ontario, Canada, or remote.
# Empty/missing locations are included (some ATS types don't return location).
ALLOWED_LOCATIONS = [
    "toronto", "ontario", "mississauga", "markham", "scarborough",
    "north york", "etobicoke", "brampton", "vaughan", "richmond hill",
    "oakville", "burlington", "hamilton", "kitchener", "waterloo",
    "canada", "remote", "hybrid", "anywhere",
]


# Locations that disqualify a job even if an allowed term is present
BLOCKED_LOCATIONS = [
    "remote us", "remote - us", "remote | us", "remote, us",
    "united states", "usa only", "us only", "us-only",
]


def is_location_allowed(job):
    """Return True if job location matches allowlist or is empty."""
    loc = job.get("location", "").strip().lower()
    if not loc:
        return True
    # Block US-specific remote roles first
    if any(term in loc for term in BLOCKED_LOCATIONS):
        return False
    return any(term in loc for term in ALLOWED_LOCATIONS)


def filter_top_jobs(jobs, max_count):
    # Hard-filter by location first, then score the rest
    local_jobs = [j for j in jobs if is_location_allowed(j)]
    _safe_print(f"  Location filter: {len(jobs)} jobs -> {len(local_jobs)} after filtering to Toronto/remote/Canada")
    scored = [(score_job(j), j) for j in local_jobs]
    scored.sort(key=lambda x: x[0], reverse=True)
    # Only include jobs with a positive score
    return [j for s, j in scored if s > 0][:max_count]


# ── Job description fetching ─────────────────────────────────────────────────

def fetch_jd_greenhouse(slug, job_id):
    url = f"https://boards-api.greenhouse.io/v1/boards/{slug}/jobs/{job_id}"
    req = urllib.request.Request(url, headers={"User-Agent": USER_AGENT})
    with urllib.request.urlopen(req, timeout=12) as resp:
        data = json.loads(resp.read())
        content_html = data.get("content", "")
        # Strip HTML tags
        return re.sub(r"<[^>]+>", " ", content_html).strip()


def fetch_jd_lever(slug, job_id):
    url = f"https://api.lever.co/v0/postings/{slug}/{job_id}"
    req = urllib.request.Request(url, headers={"User-Agent": USER_AGENT})
    with urllib.request.urlopen(req, timeout=12) as resp:
        data = json.loads(resp.read())
        parts = [data.get("descriptionBody", "")]
        for lst in data.get("lists", []):
            parts.append(lst.get("text", ""))
            for item in lst.get("content", []):
                if isinstance(item, dict):
                    parts.append(item.get("text", ""))
        return re.sub(r"<[^>]+>", " ", " ".join(parts)).strip()


def fetch_job_description(job):
    """Fetch the full job description text from the job's link."""
    link = job.get("link", "")
    try:
        # Greenhouse board API
        m = re.search(r"greenhouse\.io/(\w[\w-]*)/jobs/(\d+)", link)
        if m:
            return fetch_jd_greenhouse(m.group(1), m.group(2))

        # Lever
        m = re.search(r"jobs\.lever\.co/([\w-]+)/([\w-]+)", link)
        if m:
            return fetch_jd_lever(m.group(1), m.group(2))

        # Ashby — no public JD API; fetch page
        m = re.search(r"jobs\.ashbyhq\.com/([\w-]+)/([\w-]+)", link)
        if m:
            req = urllib.request.Request(link, headers={"User-Agent": USER_AGENT})
            with urllib.request.urlopen(req, timeout=12) as resp:
                html = resp.read().decode("utf-8", errors="replace")
                return re.sub(r"<[^>]+>", " ", html)[:4000]

    except Exception as e:
        _safe_print(f"    JD fetch failed for {job['title']}: {e}")

    # Fallback: use title + company as context
    return f"Role: {job['title']}\nCompany: {job['company']}\nLocation: {job['location']}"


# ── Resume customization ─────────────────────────────────────────────────────

def customize_resume(resume_text, jd_text, job, skill_content):
    """Customize the resume using the local 'claude -p' CLI (uses your Claude subscription)."""
    prompt = f"""You are an expert PM resume customizer using the Product Growth methodology.

ABSOLUTE RULES (violation = rejected output):
1. NEVER invent metrics, companies, dates, titles, or accomplishments not in the base resume
2. NEVER change numbers — if the base says 250+, write 250+, not 300+
3. NEVER add a company the candidate did not work at
4. You MAY: reorder, reframe, emphasize different facts, rewrite bullets to highlight relevance
5. You MAY: adjust the Summary section to speak directly to this role
6. Output: the full customized resume text only — no commentary, no preamble

RESUME CUSTOMIZATION METHODOLOGY:
{skill_content[:2000]}

JOB: {job['title']} at {job['company'].split(' (')[0]}
LOCATION: {job['location']}

JOB DESCRIPTION:
{jd_text[:3500]}

BASE RESUME (these are the ONLY facts you may use):
{resume_text}

Output the complete customized resume, ready to submit. Plain text. Keep section headers clear (SUMMARY, EXPERIENCE, EDUCATION, SKILLS)."""

    # Write prompt to a temp file to avoid Windows 8191-char command-line limit
    import tempfile
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".txt", delete=False, encoding="utf-8", errors="replace"
    ) as f:
        f.write(prompt)
        tmpfile = f.name

    try:
        # Unset CLAUDECODE so the nested session isn't blocked when run
        # from within a Claude Code session (e.g., during manual testing).
        # When Task Scheduler runs this at 6pm it won't be set anyway.
        env = os.environ.copy()
        env.pop("CLAUDECODE", None)

        with open(tmpfile, "r", encoding="utf-8", errors="replace") as fin:
            result = subprocess.run(
                CLAUDE_CMD + ["-p", "-"],
                stdin=fin,
                capture_output=True,
                text=True,
                timeout=300,
                env=env,
            )
    finally:
        os.unlink(tmpfile)

    if result.returncode != 0:
        raise RuntimeError(f"claude CLI error: {result.stderr[:300]}")
    return result.stdout.strip()


# ── Hallucination check ──────────────────────────────────────────────────────

# Facts we know are in the candidate's resume -- used to build an allowlist
_BASE_FACTS_PATTERNS = [
    r"\b250\+",
    r"\b60\s*%",
    r"\$200[Kk]",
    r"\b50\s*%",
    r"\b25\s*%",
    r"\b30\s*%",
    r"\b5[,.]000\+",
    r"\b2[,.]000\+",
    r"\b100\+",
    r"\b20\+\s*schools",
    r"\b12\s*months",
    r"\b10\+\s*countries",
    r"\b3[,.]000\+",
    r"\b\$2M\+",
    r"\b\$20M\+",
    r"\b1:450",
    r"\b1:260",
    r"\b80[,.]000\+",
    r"\b100[,.]000\+",
    r"\bBMO\b",
    r"\bUpGrad\b",
    r"\bAspiring Minds\b",
    r"\bGet Set Sorted\b",
    r"\bGoogle\b",
    r"\bSHL\b",
    r"\bFRTB\b",
    r"\bTORIC\b",
    r"\bFMS\b",
    r"\bNSIT\b",
    r"\bMississauga\b",
]


def check_hallucinations(base_text, customized_text):
    """
    Return a list of issues where the customized resume contains
    numbers or proper nouns not found in the base resume.
    """
    issues = []

    # 1. Find every % number and dollar amount in the customized text
    candidate_facts = re.findall(
        r"\b\d[\d,\.]*\s*[\+%]|\$\d[\d,\.]*[KMBkmb]?\+?", customized_text
    )
    # Allow common years-of-experience numbers (Claude infers these from timeline)
    _ALLOWED_YEAR_NUMBERS = {"7+", "8+", "9+", "10+", "11+", "12+", "13+", "14+", "15+"}
    candidate_facts = [f for f in candidate_facts if f.strip() not in _ALLOWED_YEAR_NUMBERS]
    for fact in set(candidate_facts):
        # Check that this fact appears in the base
        escaped = re.escape(fact.strip())
        if not re.search(escaped, base_text, re.IGNORECASE):
            issues.append(f"Number not in base resume: '{fact}'")

    # 2. Find multi-word proper nouns not in base (2+ title-cased words required)
    # Single capitalized words are too noisy — they catch sentence-starting words,
    # common domain terms, and month/title words. Only flag 2+ word phrases.
    multi_word_pattern = r"\b[A-Z][a-z]{2,}(?:[ \t]+[A-Z][a-z]{2,})+\b"
    custom_phrases = set(re.findall(multi_word_pattern, customized_text))
    base_phrases = set(re.findall(multi_word_pattern, base_text))
    # Words that make a phrase safe — if any word in the phrase is here, skip it
    allowed_words = {
        # Resume structure
        "Summary", "Experience", "Education", "Skills", "Tools", "Present",
        "Relevant", "Earlier", "Roles", "Profile",
        # Seniority / job titles
        "Product", "Manager", "Senior", "Staff", "Principal", "Director",
        "Head", "Lead", "Group", "Associate", "Vice", "President", "Chief",
        "Agile", "Delivery", "Architect", "Analyst", "Consultant", "Owner",
        # Business / PM domain terms
        "Financial", "Finance", "Risk", "Compliance", "Regulatory", "Audit",
        "Governance", "Reporting", "Platform", "Enterprise", "Strategy",
        "Stakeholder", "Roadmap", "Discovery", "Program", "Technical",
        "Engineering", "Operations", "Analytics", "Data", "Payments",
        "Revenue", "Growth", "Customer", "Internal", "External", "Digital",
        "Change", "Management", "Readiness", "Enablement", "Foundations",
        "Defined", "Accountable", "Responsible", "Informed", "Consulted",
        "Card", "Credit", "Fraud", "Security", "Identity", "Privacy",
        # Locations
        "Remote", "Canada", "Toronto", "Ontario", "Vancouver", "Mississauga",
        "North", "America", "India", "Delhi",
        # Time
        "January", "February", "March", "April", "May", "June",
        "July", "August", "September", "October", "November", "December",
        # Generic descriptors and action words
        "Domains", "Initiatives", "Outcomes", "Insights", "Impact",
        "Cross", "Functional", "Multi", "End", "Based", "Driven",
        "Raised", "Built", "Led", "Owned", "Shipped", "Designed",
        "Front", "Office", "Process", "Transformation", "Regulated",
        "Environments", "Environment", "Complex", "Highly", "Large",
        "Scale", "Scaled", "Real", "Time", "High", "Low", "New",
        "Key", "Core", "Full", "Global", "Modern", "Advanced",
        "Business", "Value", "Delivery", "Team", "Teams", "Work",
        "User", "Users", "People", "Service", "Services", "System",
        "Systems", "Feature", "Features", "Sprint", "Cycle", "Flow",
        # Additional PM/tech domain terms
        "Clarity", "Domain", "Lifecycle", "Portfolio", "Backlog",
        "Dashboard", "Automation", "Framework", "Pipeline",
        "Integration", "Migration", "Architecture", "Implementation",
        "Deployment", "Prioritization", "Experimentation", "Iteration",
        "Optimization", "Assessment", "Specification", "Requirements",
        "Methodology", "Infrastructure", "Orchestration", "Stakeholders",
        "Capabilities", "Enabler", "Maturity", "Adoption", "Alignment",
        "Standardization", "Consolidation", "Modernization",
    }
    base_text_lower = base_text.lower()
    new_phrases = custom_phrases - base_phrases
    for phrase in sorted(new_phrases):
        phrase_words = phrase.split()
        # Skip if any component word is a known-safe term
        if any(w in allowed_words for w in phrase_words):
            continue
        # Skip if the phrase appears in the base resume in any casing
        if phrase.lower() in base_text_lower:
            continue
        issues.append(f"Proper noun not in base resume: '{phrase}' — verify this is not a fabrication")

    return issues


# ── .docx generation ─────────────────────────────────────────────────────────

def create_resume_docx(customized_text, job, output_path):
    """Generate a formatted .docx from the customized resume text."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Remove default paragraph spacing
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    lines = [l.rstrip() for l in customized_text.split("\n")]
    is_first = True

    for line in lines:
        stripped = line.strip()

        if not stripped:
            # Small spacer paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            continue

        # Detect line type
        is_name = is_first and not stripped.startswith("-")
        is_section_header = (
            stripped.isupper() and 3 < len(stripped) < 50
            or stripped.startswith("## ")
            or stripped.startswith("# ")
        )
        is_bullet = stripped.startswith("- ") or stripped.startswith("• ")
        is_contact = any(x in stripped for x in ["@", "linkedin", "github", "+1 ", "647", "905"])

        clean = re.sub(r"^[#\-•\s]+", "", stripped).strip()

        if is_name:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(clean)
            run.bold = True
            run.font.size = Pt(16)
            p.paragraph_format.space_after = Pt(2)
            is_first = False

        elif is_contact:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(4)

        elif is_section_header:
            p = doc.add_paragraph()
            run = p.add_run(clean.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            # Underline via border
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "1a56db")
            pBdr.append(bottom)
            pPr.append(pBdr)

        elif is_bullet:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(stripped[2:])
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Inches(0.25)

        else:
            p = doc.add_paragraph()
            # Check if it looks like a job title / date header (bold it)
            if re.search(r"\b(20\d{2})\b", stripped) or stripped.endswith(":"):
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(10)
            else:
                p.add_run(stripped)
            p.paragraph_format.space_after = Pt(1)

    doc.save(output_path)
    return output_path


# ── Email building ───────────────────────────────────────────────────────────

def build_email_html(jobs_data, date_str):
    """Build the HTML body for the daily digest email."""
    job_rows = ""
    for i, item in enumerate(jobs_data, 1):
        job = item["job"]
        issues = item.get("hallucination_issues", [])
        resume_name = os.path.basename(item["resume_path"]) if item.get("resume_path") else "N/A"

        hallucination_html = ""
        if issues:
            hallucination_html = f"""
            <p style="color:#dc2626;font-size:12px;margin:4px 0 0 0;">
              ⚠️ Review before sending — {len(issues)} potential issue(s) flagged:<br>
              {"<br>".join(f"&nbsp;&nbsp;• {h}" for h in issues[:5])}
            </p>"""

        job_rows += f"""
        <tr style="border-bottom:1px solid #e5e7eb;">
          <td style="padding:16px 8px;font-size:13px;font-weight:600;color:#111827;vertical-align:top;">
            {i}
          </td>
          <td style="padding:16px 8px;vertical-align:top;">
            <p style="margin:0;font-size:14px;font-weight:700;color:#111827;">{job['title']}</p>
            <p style="margin:2px 0;font-size:13px;color:#374151;">{job['company'].split(' (')[0]}</p>
            <p style="margin:2px 0;font-size:12px;color:#6b7280;">📍 {job['location']}</p>
            {hallucination_html}
          </td>
          <td style="padding:16px 8px;vertical-align:top;text-align:center;">
            <a href="{job['link']}" style="display:inline-block;background:#1d4ed8;color:white;padding:6px 14px;border-radius:4px;text-decoration:none;font-size:12px;font-weight:600;">Apply</a>
          </td>
          <td style="padding:16px 8px;vertical-align:top;font-size:11px;color:#6b7280;">
            📎 {resume_name}
          </td>
        </tr>"""

    return f"""<!DOCTYPE html>
<html>
<body style="font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;background:#f9fafb;margin:0;padding:20px;">
  <div style="max-width:700px;margin:0 auto;background:white;border-radius:8px;overflow:hidden;box-shadow:0 1px 3px rgba(0,0,0,0.1);">

    <div style="background:#1d4ed8;padding:24px 28px;">
      <h1 style="color:white;margin:0;font-size:20px;">PM Job Digest</h1>
      <p style="color:#bfdbfe;margin:4px 0 0 0;font-size:13px;">{date_str} · Top {len(jobs_data)} new matches · Customized resumes attached</p>
    </div>

    <div style="padding:20px 28px;">
      <p style="color:#374151;font-size:13px;margin:0 0 16px 0;">
        These are today's highest-signal new PM roles based on your background in risk tech, fintech, and data platforms.
        Each attachment is a resume customized for that specific role.
        <strong>Review hallucination warnings before applying.</strong>
      </p>

      <table style="width:100%;border-collapse:collapse;">
        <thead>
          <tr style="background:#f3f4f6;">
            <th style="padding:8px;text-align:left;font-size:11px;color:#6b7280;width:28px;">#</th>
            <th style="padding:8px;text-align:left;font-size:11px;color:#6b7280;">Role</th>
            <th style="padding:8px;text-align:center;font-size:11px;color:#6b7280;width:80px;">Link</th>
            <th style="padding:8px;text-align:left;font-size:11px;color:#6b7280;">Resume</th>
          </tr>
        </thead>
        <tbody>
          {job_rows}
        </tbody>
      </table>
    </div>

    <div style="background:#f3f4f6;padding:16px 28px;">
      <p style="margin:0;font-size:11px;color:#9ca3af;">
        Generated by PM Job Fetcher · Customized with Claude API · Verify all facts before submitting
      </p>
    </div>

  </div>
</body>
</html>"""


# ── Email sending ─────────────────────────────────────────────────────────────

def send_email(jobs_data, settings, date_str):
    """Send the digest email with .docx attachments via Gmail SMTP."""
    msg = MIMEMultipart()
    msg["From"] = settings["email_from"]
    msg["To"] = settings["email_to"]
    msg["Subject"] = (
        f"PM Job Digest {date_str} — {len(jobs_data)} new match{'es' if len(jobs_data) != 1 else ''}"
    )

    html = build_email_html(jobs_data, date_str)
    msg.attach(MIMEText(html, "html"))

    for item in jobs_data:
        path = item.get("resume_path")
        if path and os.path.exists(path):
            with open(path, "rb") as f:
                attach = MIMEApplication(f.read(), _subtype="docx")
                attach["Content-Disposition"] = (
                    f'attachment; filename="{os.path.basename(path)}"'
                )
                msg.attach(attach)

    print(f"Sending email to {settings['email_to']}...", end=" ", flush=True)
    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
        smtp.login(settings["email_from"], settings["email_app_password"])
        smtp.send_message(msg)
    print("sent!")


# ── WhatsApp sending ─────────────────────────────────────────────────────────

WHATSAPP_API_BASE = "https://graph.facebook.com/v21.0/{phone_id}"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _whatsapp_upload_media(phone_id, token, file_path):
    """Upload a file to WhatsApp media and return the media_id."""
    url = f"{WHATSAPP_API_BASE.format(phone_id=phone_id)}/media"

    # Build multipart/form-data manually (no requests library)
    boundary = "----WhatsAppMediaBoundary"
    filename = os.path.basename(file_path)

    with open(file_path, "rb") as f:
        file_data = f.read()

    body = (
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="messaging_product"\r\n\r\n'
        f"whatsapp\r\n"
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="file"; filename="{filename}"\r\n'
        f"Content-Type: {DOCX_MIME}\r\n\r\n"
    ).encode("utf-8") + file_data + f"\r\n--{boundary}--\r\n".encode("utf-8")

    headers = {
        "Content-Type": f"multipart/form-data; boundary={boundary}",
        "Authorization": f"Bearer {token}",
    }

    req = urllib.request.Request(url, data=body, headers=headers, method="POST")
    with urllib.request.urlopen(req) as resp:
        result = json.loads(resp.read().decode("utf-8"))
    return result["id"]


def _whatsapp_send_document(phone_id, token, to_number, media_id, filename, caption):
    """Send a document message via WhatsApp using an uploaded media_id."""
    url = f"{WHATSAPP_API_BASE.format(phone_id=phone_id)}/messages"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {token}",
    }
    payload = json.dumps({
        "messaging_product": "whatsapp",
        "recipient_type": "individual",
        "to": to_number,
        "type": "document",
        "document": {
            "id": media_id,
            "filename": filename,
            "caption": caption,
        },
    })
    req = urllib.request.Request(url, data=payload.encode("utf-8"),
                                 headers=headers, method="POST")
    with urllib.request.urlopen(req) as resp:
        resp.read()


def send_whatsapp(jobs_data, settings):
    """Send one WhatsApp template message per job, with resume attached."""
    phone_id = settings.get("whatsapp_phone_id")
    token = settings.get("whatsapp_token")
    to_number = settings.get("whatsapp_to")
    template_name = settings.get("whatsapp_template", "job_alerts")

    if not all([phone_id, token, to_number]):
        print("WhatsApp not configured (missing phone_id/token/to). Skipping.")
        return

    msg_url = f"{WHATSAPP_API_BASE.format(phone_id=phone_id)}/messages"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {token}",
    }

    sent = 0
    failed = 0
    for item in jobs_data:
        job = item["job"]
        company = job["company"].split(" (")[0]
        resume_path = item.get("resume_path")

        # 1. Send template message (job alert)
        payload = json.dumps({
            "messaging_product": "whatsapp",
            "recipient_type": "individual",
            "to": to_number,
            "type": "template",
            "template": {
                "name": template_name,
                "language": {"code": "en"},
                "components": [
                    {
                        "type": "body",
                        "parameters": [
                            {"type": "text", "text": job["title"]},
                            {"type": "text", "text": company},
                            {"type": "text", "text": job["location"]},
                            {"type": "text", "text": job["link"]},
                        ],
                    }
                ],
            },
        })

        try:
            req = urllib.request.Request(msg_url, data=payload.encode("utf-8"),
                                         headers=headers, method="POST")
            with urllib.request.urlopen(req) as resp:
                resp.read()
            sent += 1

            # 2. Attach resume if available
            if resume_path and os.path.exists(resume_path):
                try:
                    media_id = _whatsapp_upload_media(phone_id, token, resume_path)
                    _whatsapp_send_document(
                        phone_id, token, to_number, media_id,
                        filename=os.path.basename(resume_path),
                        caption=f"Customized resume for {job['title']} at {company}",
                    )
                except Exception as e:
                    _safe_print(f"  WhatsApp resume attach failed for {company}: {e}")

        except urllib.error.HTTPError as e:
            error_body = e.read().decode("utf-8", errors="replace")
            _safe_print(f"  WhatsApp failed for {company}: {e.code} {error_body}")
            failed += 1
        except Exception as e:
            _safe_print(f"  WhatsApp failed for {company}: {e}")
            failed += 1

    print(f"WhatsApp: {sent} sent, {failed} failed")


# ── Job trackers ─────────────────────────────────────────────────────────────

MARKDOWN_TRACKER = os.path.join(SCRIPT_DIR, "job_tracker.md")


def update_markdown_tracker(jobs_data, date):
    """Prepend today's matches to job_tracker.md."""
    date_str = date.strftime("%B %d, %Y")
    lines = [f"\n## {date_str}\n\n"]
    lines.append("| Status | Role | Company | Location | Link |\n")
    lines.append("|--------|------|---------|----------|------|\n")
    for item in jobs_data:
        j = item["job"]
        company = j["company"].split(" (")[0]
        lines.append(
            f"| [ ] | {j['title']} | {company} | {j['location']} | [Apply]({j['link']}) |\n"
        )
    lines.append("\n---\n")

    # Read existing content after the header
    header = ""
    body = ""
    if os.path.exists(MARKDOWN_TRACKER):
        with open(MARKDOWN_TRACKER, encoding="utf-8") as f:
            content = f.read()
        # Split at first "---"
        parts = content.split("---\n", 1)
        header = parts[0] + "---\n"
        body = parts[1] if len(parts) > 1 else ""
    else:
        header = (
            "# PM Job Tracker\n\n"
            "Running log of top-matched PM roles, newest first.\n"
            "Status: `[ ]` not applied · `[→]` applied · `[✓]` heard back · `[✗]` passed/rejected\n\n"
            "---\n"
        )

    new_content = header + "".join(lines) + body
    with open(MARKDOWN_TRACKER, "w", encoding="utf-8") as f:
        f.write(new_content)
    print(f"  Markdown tracker updated: job_tracker.md")


def update_excel_tracker(jobs_data, date, settings):
    """Append today's top matches to the Applications tab of the Excel tracker."""
    try:
        import openpyxl
        from openpyxl.styles import PatternFill, Font, Alignment
    except ImportError:
        print("  Skipping Excel update (openpyxl not installed)")
        return

    excel_path = settings.get("excel_tracker_path", "")
    if not os.path.exists(excel_path):
        print(f"  Excel tracker not found at: {excel_path}")
        return

    try:
        wb = openpyxl.load_workbook(excel_path)
        ws = wb["Applications"]

        # Find the first truly empty row (max_row includes blank formatted rows)
        next_row = 2
        for row in ws.iter_rows(min_row=2, max_col=2):
            if any(cell.value for cell in row):
                next_row = row[0].row + 1

        # Light yellow fill for new rows so they stand out
        yellow_fill = PatternFill(start_color="FFFDE7", end_color="FFFDE7", fill_type="solid")

        # Collect existing (company, job) pairs to avoid duplicates
        existing = set()
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[1] and row[2]:
                existing.add((str(row[1]).strip().lower(), str(row[2]).strip().lower()))

        added = 0
        for item in jobs_data:
            j = item["job"]
            company = j["company"].split(" (")[0]
            title_with_link = f"{j['title']}  →  {j['link']}"
            key = (company.lower(), j["title"].lower())

            if key in existing:
                continue  # already tracked

            row_data = [
                date.strftime("%Y-%m-%d"),  # A: Date
                company,                     # B: Company
                title_with_link,             # C: Job (title + link)
                "",                          # D: Application Type
                "",                          # E: Next Steps
                "New",                       # F: Status
                date.strftime("%Y-%m-%d"),   # G: Last Updated
            ]
            for col, value in enumerate(row_data, start=1):
                cell = ws.cell(row=next_row, column=col, value=value)
                cell.fill = yellow_fill
                cell.alignment = Alignment(wrap_text=False)

            existing.add(key)
            next_row += 1
            added += 1

        wb.save(excel_path)
        print(f"  Excel tracker updated: {added} new row(s) added to Applications tab")

    except Exception as e:
        print(f"  Excel update failed: {e}")


# ── Main orchestration ────────────────────────────────────────────────────────

def main():
    test_mode = "--test" in sys.argv
    use_today = "--today" in sys.argv or "--test" in sys.argv

    print("=" * 60)
    print("PM Job Digest")
    print("=" * 60)

    settings = load_settings()
    validate_settings(settings, test_mode=test_mode)

    # Step 1: Fetch new jobs (unless using existing output)
    if not use_today:
        print("\nFetching new PM jobs...")
        result = subprocess.run(
            ["python", os.path.join(SCRIPT_DIR, "fetch_jobs.py")],
            cwd=SCRIPT_DIR,
            capture_output=False,
        )
        if result.returncode != 0:
            print("WARNING: fetch_jobs.py returned non-zero exit code")

    # Step 2: Load latest output
    output_file = find_latest_output()
    if not output_file:
        print("No output file found. Run fetch_jobs.py first.")
        sys.exit(1)

    today = datetime.date.today()
    date_str = today.strftime("%B %d, %Y")
    print(f"\nUsing: {os.path.basename(output_file)}")

    all_jobs = parse_jobs_from_file(output_file)
    print(f"Parsed {len(all_jobs)} jobs")

    # Step 3: Score and filter
    max_jobs = settings.get("max_jobs_per_digest", 5)
    top_jobs = filter_top_jobs(all_jobs, max_jobs)

    if not top_jobs:
        print("No new matching jobs today. No email sent.")
        return

    print(f"\nTop {len(top_jobs)} matches:")
    for j in top_jobs:
        _safe_print(f"  [{score_job(j):2d}] {j['company'].split(' (')[0]}: {j['title']} -- {j['location']}")

    # Step 4: Load base resume and skill
    print("\nLoading base resume and skill...")
    base_resume_text = load_base_resume(settings)
    skill_content = load_skill_content(settings)

    if not base_resume_text:
        print("ERROR: Could not load base resume. Check base_resume_path in settings.json")
        sys.exit(1)

    # Step 5: For each job — fetch JD, customize, check, generate .docx
    os.makedirs(RESUMES_DIR, exist_ok=True)
    today_slug = today.strftime("%Y-%m-%d")

    jobs_data = []
    for i, job in enumerate(top_jobs, 1):
        company_clean = re.sub(r"[^\w\s-]", "", job["company"].split(" (")[0]).strip()[:25]
        title_clean = re.sub(r"[^\w\s-]", "", job["title"]).strip()[:45]
        # Human-readable name: "Affirm - Senior PM Financial Reporting.docx"
        resume_filename = f"{company_clean} - {title_clean}.docx"
        resume_path = os.path.join(RESUMES_DIR, resume_filename)

        _safe_print(f"\n[{i}/{len(top_jobs)}] {job['title']} @ {job['company'].split(' (')[0]}")

        # Fetch JD
        print(f"  Fetching job description...", end=" ", flush=True)
        jd_text = fetch_job_description(job)
        print(f"got {len(jd_text)} chars")

        # Customize resume
        customized_text = None
        hallucination_issues = []

        if not test_mode:
            print(f"  Customizing resume via Claude CLI...", end=" ", flush=True)
            try:
                customized_text = customize_resume(
                    base_resume_text, jd_text, job, skill_content
                )
                print("done")

                # Hallucination check
                print(f"  Checking for hallucinations...", end=" ", flush=True)
                hallucination_issues = check_hallucinations(base_resume_text, customized_text)
                if hallucination_issues:
                    _safe_print(f"  WARNING: {len(hallucination_issues)} issue(s) flagged")
                    for issue in hallucination_issues[:3]:
                        _safe_print(f"    - {issue}")
                else:
                    _safe_print("  clean")

            except Exception as e:
                print(f"ERROR: {e}")
                customized_text = None

        # Generate .docx
        if customized_text:
            print(f"  Generating .docx...", end=" ", flush=True)
            try:
                create_resume_docx(customized_text, job, resume_path)
                print(f"saved: {resume_filename}")
            except Exception as e:
                print(f"ERROR generating docx: {e}")
                resume_path = None
        else:
            resume_path = None

        jobs_data.append({
            "job": job,
            "jd_text": jd_text,
            "customized_text": customized_text,
            "hallucination_issues": hallucination_issues,
            "resume_path": resume_path,
        })

    # Step 6: Send email + WhatsApp
    if test_mode:
        print(f"\n[TEST MODE] Skipping email and WhatsApp send.")
        print(f"Would have sent to: {settings.get('email_to', 'not set')}")
        print(f"Would have sent {len(jobs_data)} WhatsApp message(s) to: {settings.get('whatsapp_to', 'not set')}")
        print(f"Subject: PM Job Digest {date_str} — {len(jobs_data)} new matches")
        print("\nHTML preview saved to: digest_preview.html")
        with open(os.path.join(SCRIPT_DIR, "digest_preview.html"), "w", encoding="utf-8") as f:
            f.write(build_email_html(jobs_data, date_str))
    else:
        send_email(jobs_data, settings, date_str)
        send_whatsapp(jobs_data, settings)
        print("\nDone! Check your inbox and WhatsApp.")

    # Always update the trackers (markdown + Excel)
    update_markdown_tracker(jobs_data, today)
    update_excel_tracker(jobs_data, today, settings)


if __name__ == "__main__":
    main()
